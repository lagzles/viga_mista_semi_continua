from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt

import vm_db as db
import vm_db_gerais as db_gerais
import vm_db_relatorio as db_relatorio


dbt = r'vmsc.db'
dbg = r'vmsc_g.db'
dbrel = r'vmsc_relatorio.db'

def relatorio():
    try:
        document = Document()
        document.add_picture('medabil.jpg', width=Inches(2.5))

    ##    document.add_heading('Relatório VMSC', 0)
        ## Estilo para titulo
        obj_styles = document.styles
        obj_charstyle = obj_styles.add_style('sha', WD_STYLE_TYPE.CHARACTER)
        obj_font = obj_charstyle.font
        obj_font.size = Pt(48)
        obj_font.name = 'Times New Roman'
        # Fim do estilo do titulo

        pt = document.add_paragraph()
        pt.add_run().add_break(WD_BREAK.LINE)
        pt.add_run().add_break(WD_BREAK.LINE)
        pt.add_run().add_break(WD_BREAK.LINE)
        pt.add_run().add_break(WD_BREAK.LINE)   

        style = document.styles['Strong']
            
        style.font.size = Pt(48)
        style.font.name = 'Calibri'
        pt.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        pt.add_run("Relatório de Cálculo do Sistema de Viga Mista Semi Continua", style='Strong')
        
        document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        paragraph = document.add_paragraph()
        run = paragraph.add_run() 
        
        fldChar = OxmlElement('w:fldChar')  # creates a new element
        fldChar.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'   # change 1-3 depending on heading levels you need

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        fldChar3 = OxmlElement('w:t')
        fldChar3.text = "Right-click to update field."
        fldChar2.append(fldChar3)

        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(qn('w:fldCharType'), 'end')

        r_element = run._r
        r_element.append(fldChar)
        r_element.append(instrText)
        r_element.append(fldChar2)
        r_element.append(fldChar4)
        p_element = paragraph._p

        pbreak0 = document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    ##    run.add_break(WD_BREAK.PAGE)
        document.add_heading('Introdução', level=1)
        p = document.add_paragraph("O presente relatorio tem por proposito apresentar as premissas de calculo e um resumo da rotina de calculo realizado, apresentando os resultados obtidos.")
        p = document.add_paragraph("O Relatorio apresentara primeiramente as considerações iniciais, assim como as propriedades adotadas na rotina de calculo.")
        p = document.add_paragraph("Por fim apresentara um resumo dos resultados obtidos no final das rotinas de calculo.")
    ##    p.page_break_after = True

        pbreak = document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

        ## Cria tabela de Propriedades dos materiais
        document.add_heading('Considerações iniciais de Cálculo:', level=1)
        document.add_heading('Propriedades dos Materiais', level=2)
        p = document.add_paragraph("Tabela com valores de resistência dos materiais utilizados na rotina de escoamento.")
        p.add_run("Sendo fy a tensão de escoamento do aço do perfil metálico,")
        p.add_run("fys a tensão de escoamento do aço das barras de armadura da laje e ")
        p.add_run("fck a resistência caracteristica do concreto.")
        run = p.add_run()
        run.style = None
        
        table = document.add_table(rows=3, cols=4)
        table.style = 'Table Grid'
        
        table.cell(0, 0).text = "fy  [kN/cm2] ="
        table.cell(1, 0).text = "fs  [kN/cm2] ="
        table.cell(2, 0).text = "fck [kN/cm2] ="
        
        table.cell(0, 1).text = str(db_gerais.get_data(dbg, 'fy'))
        table.cell(1, 1).text = str(db_gerais.get_data(dbg, 'fys'))
        table.cell(2, 1).text = str(db_gerais.get_data(dbg, 'fck'))
        
        table.cell(0, 2).text = "E   [kN/m2] ="
        table.cell(1, 2).text = "Es  [kN/m2] ="
        table.cell(2, 2).text = "G   [kN/m2] ="
        
        table.cell(0, 3).text = str(db_gerais.get_data(dbg, 'e'))
        table.cell(1, 3).text = str(db_gerais.get_data(dbg, 'es'))
        table.cell(2, 3).text = str(db_gerais.get_data(dbg, 'G'))

        p = document.add_paragraph(" ")
        
        ## Cria tabela de Propriedades de ligação, valores que não podem ser
        ## alterados pelo usuario
        document.add_heading('Propriedades das Ligações', level=2)
        p = document.add_paragraph("Tabela com valores padrão para propriedades das ligações, utilizados nos calculos de verificações e vinculos semi-rigidos.")
        table = document.add_table(rows=8, cols=3)    
        table.style = 'Table Grid'

        table.cell(0, 1).text = "Quantidade"
        table.cell(0, 2).text = "Unidade"
        
        table.cell(1, 0).text = "Espessura Cantoneira de alma"
        table.cell(2, 0).text = "Espessura Cantoneiras inferiores"
        table.cell(3, 0).text = "Resistência Aço das cantoneiras"
        table.cell(4, 0).text = "Diâmetro dos parafusos"
        table.cell(5, 0).text = "Resistência Aço dos parafusos"
        table.cell(6, 0).text = "Grau de interação"
        table.cell(7, 0).text = "Resistência dos Conectores"

        table.cell(1, 1).text = "6.35"
        table.cell(2, 1).text = "9.50"
        table.cell(3, 1).text = "46.0"
        table.cell(4, 1).text = "19.05"
        table.cell(5, 1).text = "82.5"
        table.cell(6, 1).text = "0.75"
        table.cell(7, 1).text = "70.7"

        table.cell(1, 2).text = "mm"
        table.cell(2, 2).text = "mm"
        table.cell(3, 2).text = "kN/cm2"
        table.cell(4, 2).text = "mm"
        table.cell(5, 2).text = "kN/cm2"
        table.cell(6, 2).text = "%"
        table.cell(7, 2).text = "kN"
        p = document.add_paragraph(" ")

        ## Cria tabela de cargas consideradas
        document.add_heading('Carregamentos e Combinações Consideradas', level=2)
        
        p = document.add_paragraph("Tabela apresentando os valores dos carregamentos considerados em calculo.")
        p.add_run("Considerando um steel deck de espessura {}mm e laje com altura de {}mm.".format(db_gerais.get_data(dbg,'deck'),db_gerais.get_data(dbg,'hdeck')))
        table = document.add_table(rows=5, cols=4)
        table.style = 'Table Grid'

        table.cell(0, 1).text = "Quantidade"
        table.cell(0, 2).text = "Unidade"
        table.cell(0, 3).text = "Sigla"
        
        table.cell(1, 0).text = "Carga Steel Deck"
        table.cell(2, 0).text = "Carga Permanente"
        table.cell(3, 0).text = "Sobrecarga"
        table.cell(4, 0).text = "Sobrecarga de Montagem"

        table.cell(1, 2).text = "kg/m2"
        table.cell(2, 2).text = "kg/m2"
        table.cell(3, 2).text = "kg/m2"
        table.cell(4, 2).text = "kg/m2"

        table.cell(1, 3).text = "cpsd"
        table.cell(2, 3).text = "cp"
        table.cell(3, 3).text = "sc"
        table.cell(4, 3).text = "scm"
        
        table.cell(1, 1).text = str(db_gerais.get_data(dbg,'q_cpsd'))  # str(250.)
        table.cell(2, 1).text = str(db_gerais.get_data(dbg,'q_cp'))  # str(200.)
        table.cell(3, 1).text = str(db_gerais.get_data(dbg,'q_sc'))  # str(300.)
        table.cell(4, 1).text = str(100.)
        
        pbreak22 = document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        
        document.add_heading("Combinações utilizadas", level=2)    
        document.add_paragraph('ML  \t 1.0 * sc + cp  \t', style='List Number')
        document.add_paragraph('MD\' \t 1.2 * (cpsd + pp) + 1.6 * scm \t (Antes da Cura)', style='List Number')
        document.add_paragraph('MD  \t 1.2 * (cpsd + pp) + 1.6 * sc \t (Depois da Cura)', style='List Number')
        document.add_paragraph('Longa \t cp + cpsd + pp', style='List Number')
        document.add_paragraph('Curta \t sc', style='List Number')
        p = document.add_paragraph(" ")
        

        # Cria tabela de Geometria do Sistema de Tramos
        document.add_heading('Propriedades Geométricas', level=2)
        p = document.add_paragraph("Tabela apresentando as dimensões geométricas dos perfis que compõe cada vão.")
        
        n =  db.count_data(dbt) # 5 #len(tramos)
        table = document.add_table(rows=9, cols=n+1)    
        table.style = 'Table Grid'
        
        table.cell(0, 0).text = "Indice do Vão "
        table.cell(1, 0).text = "Vão [m]"
        table.cell(2, 0).text = "b inf [m]"
        table.cell(3, 0).text = "alma [mm]"
        table.cell(4, 0).text = "mesa [mm]"
        table.cell(5, 0).text = "tw [mm]"
        table.cell(6, 0).text = "tfs [mm]"
        table.cell(7, 0).text = "tfi [mm]"
        table.cell(8, 0).text = "As"

        for i in range(n):
            table.cell(0, i+1).text = str(db.get_data(dbt, i)[0])
            table.cell(1, i+1).text = str(db.get_data(dbt, i)[1])
            table.cell(2, i+1).text = str(db.get_data(dbt, i)[2])
            table.cell(3, i+1).text = str(db.get_data(dbt, i)[3])
            table.cell(4, i+1).text = str(db.get_data(dbt, i)[4])
            table.cell(5, i+1).text = str(db.get_data(dbt, i)[5])
            table.cell(6, i+1).text = str(db.get_data(dbt, i)[6])
            table.cell(7, i+1).text = str(db.get_data(dbt, i)[7])
            fi = float(db.get_data(dbt, i)[9])
            asl = fi**2*3.14/4 * float(db.get_data(dbt, i)[8])
            table.cell(8, i+1).text = str(round(asl,2))  # "4 #12.5"

        pbreak2 = document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        

        # Calculo das propriedades das ligações semi-rigidas
        document.add_heading('Calculo das Ligações semi-rigidas', level=1)
        p = document.add_paragraph("Esta parte do relatorio ira apresentar os calculos realizados,")
        p.add_run(" para cada tramo, para o calculo do vinculo semi-rigido e outras propriedades semi-rigidas.")
        apoios = ["esquerdo","direito"]
        
        for i in range(n):
            for j in apoios:
                document.add_heading('Tramo {} apoio {}'.format(i, j), level=2)
                p = document.add_paragraph(db_relatorio.get_data(dbrel, 'Calculo do C', (str(i)+j)))

        pbreak3 = document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

        # Calculo das propriedades Equivalentes de Curta Duração
        document.add_heading('Propriedades Equivalentes de Curta Duração', level=0)
        document.add_heading('Calculo das Propriedades Equivalentes de Curta Duração', level=1)
        p = document.add_paragraph("Esta parte do relatorio ira apresentar os calculos realizados, ")
        p.add_run("para o calculo das propriedades equivalentes de cada tramo, para o caso de carregamento ")
        p.add_run("de curta duração.")
        
        for i in range(n):
            document.add_heading('Tramo {}'.format(i), level=2)
            p = document.add_paragraph(db_relatorio.get_data(dbrel, 'Equivalente Curta', i))

        pbreak4 = document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

        # Calculo das propriedades Equivalente de Longa Duração
        document.add_heading('Propriedades Equivalentes de Longa Duração', level=0)
        document.add_heading('Calculo das Propriedades Equivalentes de Longa Duração', level=1)
        p = document.add_paragraph("Esta parte do relatorio ira apresentar os calculos realizados,")
        p.add_run("para o calculo das propriedades equivalentes de cada tramo, para o caso de carregamento ")
        p.add_run("de longa duração.")

        for i in range(n):
            document.add_heading('Tramo {}'.format(i), level=2)
            p = document.add_paragraph(db_relatorio.get_data(dbrel, 'Equivalente Longa', i))


        pbreak3 = document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

        # Calculo das verificações
        document.add_heading('Verificações', level=0)
        document.add_heading('Verificações Realizadas', level=1)
        p = document.add_paragraph("Esta parte do relatorio ira apresentar os calculos realizados,")
        p.add_run("para o calculo das verificações de dimensionamento, apresentadas a seguir:")
        p = document.add_paragraph("""        
Informações sobre as verificações realizadas:
    - 1 Verificação: Cargas antes da cura
    - 2 Verificação: Esbeltez da seção
                     (NBR8800/2008 O.1.1.2.d)
    - 3 Verificação: Momento Positivo Reduzido
                     (NBR8800/2008 O.2.3)
    - 4 Verificação: Cortante
                     (NBR8800/2008 O.3)
    - 5 Verificação: Capacidade de Rotação x Rotação Necessária
                     (NBR8800/2008 Tabela R.3)
    - 6 Verificação: Flambagem lateral com distorção da
      Seção Transversal (NBR8800/2008 O.2.5)
    - 7 Verificação: Numero de Conectores
                     (NBR8800/2008 O.2.4.3)
    - 8 Verificação: Cisalhamento Longitudinal da Laje
                     (NBR8800/2008 O.1.3.4)
    - 9 Verificação: Relação entre Momento Resistivo Negativo/Positivo
    - 10 Verificação: Limitação das tensões de serviço
      (combinação rara de ações)
    - 11 Verificação: Flecha/Deslocamentos
    - 12 Verificação: Estado Limite de Vibração Excessiva
    - 13 Verificação: Fissuração do Concreto sobre os Apoios""")
    
        document.add_heading('Calculo das Verificações ', level=1)
        
        for i in range(n):
            document.add_heading('Tramo {}'.format(i), level=2)
            p = document.add_paragraph(db_relatorio.get_data(dbrel, "Verificacao", str(i)))


        pbreak3 = document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

        document.save('Relatório_VMSC.docx')
        return True
    except Exception as e:
        print(e)
        return False


##relatorio()
